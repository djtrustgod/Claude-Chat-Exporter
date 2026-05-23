"""Render a Chat to a Microsoft Word .docx file."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from ..core import markdown_walker as mdw
from ..core.models import (
    Attachment,
    Chat,
    ContentPart,
    ImagePart,
    Message,
    TextPart,
    ToolResultPart,
    ToolUsePart,
)
from .base import (
    AssetStager,
    ExportResult,
    Exporter,
    format_timestamp,
    role_label,
    safe_slug,
)


_PAGE_WIDTH_INCHES = 6.0   # Letter (8.5") minus 1.25" margins each side
_MAX_IMAGE_WIDTH_INCHES = 5.5


class DocxExporter(Exporter):
    extension = "docx"

    def export(self, chat: Chat, work_dir: Path, stager: AssetStager) -> ExportResult:
        slug = safe_slug(chat.display_name)
        doc_path = work_dir / f"{slug}.docx"

        document = Document()
        self._configure_styles(document)
        self._write_header(document, chat)

        for msg in chat.messages:
            self._write_message(document, msg, stager)

        document.save(str(doc_path))
        return ExportResult(
            work_dir=work_dir,
            doc_path=doc_path,
            attachments_dir=stager.attachments_dir,
            slug=slug,
        )

    # ------------------------------------------------------------------
    def _configure_styles(self, document: DocumentT) -> None:
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        for section in document.sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)

    def _write_header(self, document: DocumentT, chat: Chat) -> None:
        document.add_heading(chat.display_name, level=0)
        p = document.add_paragraph()
        if chat.created_at:
            self._meta_run(p, "Created: ", chat.created_at.isoformat())
        if chat.updated_at:
            self._meta_run(p, "  Updated: ", chat.updated_at.isoformat())
        if chat.model:
            self._meta_run(p, "  Model: ", chat.model)
        if chat.uuid:
            self._meta_run(p, "  ID: ", chat.uuid)
        document.add_paragraph()

    def _meta_run(self, paragraph, label: str, value: str) -> None:
        r1 = paragraph.add_run(label)
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = paragraph.add_run(value)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ------------------------------------------------------------------
    def _write_message(
        self, document: DocumentT, msg: Message, stager: AssetStager
    ) -> None:
        # Role header
        head = document.add_paragraph()
        head.paragraph_format.space_before = Pt(12)
        head.paragraph_format.space_after = Pt(2)
        r = head.add_run(role_label(msg))
        r.bold = True
        r.font.size = Pt(13)
        if msg.sender == "human":
            r.font.color.rgb = RGBColor(0x16, 0x4A, 0x84)
        elif msg.sender == "assistant":
            r.font.color.rgb = RGBColor(0x8B, 0x33, 0x0E)
        ts = format_timestamp(msg)
        if ts:
            r2 = head.add_run(f"  ·  {ts}")
            r2.font.size = Pt(9)
            r2.italic = True
            r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        for part in msg.parts:
            self._write_part(document, part, stager)

        if msg.attachments:
            self._write_attachments(document, msg.attachments, stager)

        # Light horizontal divider via a thin grey paragraph border
        divider = document.add_paragraph()
        divider.paragraph_format.space_before = Pt(6)
        divider.paragraph_format.space_after = Pt(6)
        _add_bottom_border(divider, color="DDDDDD", size_eighth_pts=4)

    # ------------------------------------------------------------------
    def _write_part(
        self, document: DocumentT, part: ContentPart, stager: AssetStager
    ) -> None:
        if isinstance(part, TextPart):
            self._render_markdown(document, part.markdown, stager)
        elif isinstance(part, ImagePart):
            self._render_image_part(document, part, stager)
        elif isinstance(part, ToolUsePart):
            self._render_code_block(
                document, f"Tool call: {part.name}", part.raw_input, lang="json"
            )
        elif isinstance(part, ToolResultPart):
            self._render_quote(document, "Tool result", part.output)

    def _render_image_part(
        self, document: DocumentT, part: ImagePart, stager: AssetStager
    ) -> None:
        staged = stager.lookup_image(part)
        if staged is None:
            p = document.add_paragraph()
            r = p.add_run(f"[Missing image: {part.filename or 'image'}]")
            r.italic = True
            r.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
            return
        p = document.add_paragraph()
        p.alignment = 1  # center
        run = p.add_run()
        width = _fit_image_width(staged.abs_path, _MAX_IMAGE_WIDTH_INCHES)
        try:
            run.add_picture(str(staged.abs_path), width=Inches(width))
        except Exception:  # noqa: BLE001
            r = p.add_run(f"[Could not embed image: {staged.abs_path.name}]")
            r.italic = True
        if staged.alt:
            cap = document.add_paragraph()
            cap.alignment = 1
            cap_run = cap.add_run(staged.alt)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _render_code_block(
        self, document: DocumentT, title: str | None, source: str, lang: str = ""
    ) -> None:
        if title:
            p = document.add_paragraph()
            r = p.add_run(title + (f"  ({lang})" if lang else ""))
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Use a single-cell table so we get a real shaded background.
        table = document.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade_cell(cell, "F4F4F4")
        cell.paragraphs[0].clear()
        for i, line in enumerate(source.splitlines() or [""]):
            para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(line)
            run.font.name = "Consolas"
            _set_complex_font(run, "Consolas")
            run.font.size = Pt(9.5)

    def _render_quote(self, document: DocumentT, title: str | None, text: str) -> None:
        if title:
            p = document.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for line in text.splitlines() or [""]:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _add_left_border(p, color="888888", size_eighth_pts=12)
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ------------------------------------------------------------------
    # Markdown → DOCX
    # ------------------------------------------------------------------
    def _render_markdown(
        self, document: DocumentT, markdown_text: str, stager: AssetStager
    ) -> None:
        state = _DocxState(document)
        for ev in mdw.walk(markdown_text):
            self._dispatch_event(ev, state)
        # In case the document ends mid-block (defensive — should not happen).
        state.finish()

    def _dispatch_event(self, ev, state: "_DocxState") -> None:
        if isinstance(ev, mdw.BlockOpen):
            self._open_block(ev, state)
        elif isinstance(ev, mdw.BlockClose):
            self._close_block(ev, state)
        elif isinstance(ev, mdw.Run):
            state.add_run(ev)
        elif isinstance(ev, mdw.InlineImage):
            # Markdown-embedded image references — usually we don't have a
            # staged file for these because images come through ImagePart.
            # Render as a placeholder link.
            state.add_run(
                mdw.Run(f"[image: {ev.alt or ev.src}]", italic=True, href=ev.src or None)
            )
        elif isinstance(ev, mdw.HardBreak):
            state.hard_break()

    def _open_block(self, ev: mdw.BlockOpen, state: "_DocxState") -> None:
        kind = ev.kind
        if kind == "paragraph":
            state.begin_paragraph()
        elif kind == "heading":
            state.begin_heading(ev.level)
        elif kind == "code_block":
            # Self-closing in our walker — render now.
            self._render_code_block(state.document, None, ev.source, lang=ev.lang)
        elif kind == "list":
            state.list_stack.append(("ordered" if ev.ordered else "bullet", 0))
        elif kind == "list_item":
            state.begin_list_item()
        elif kind == "blockquote":
            state.in_blockquote_depth += 1
        elif kind == "table":
            state.begin_table()
        elif kind in ("thead", "tbody"):
            state.table_section = kind
        elif kind == "tr":
            state.begin_table_row()
        elif kind in ("th", "td"):
            state.begin_table_cell(kind == "th")
        elif kind == "hr":
            p = state.document.add_paragraph()
            _add_bottom_border(p, color="888888", size_eighth_pts=6)

    def _close_block(self, ev: mdw.BlockClose, state: "_DocxState") -> None:
        kind = ev.kind
        if kind == "paragraph":
            state.end_paragraph()
        elif kind == "heading":
            state.end_paragraph()
        elif kind == "code_block":
            pass
        elif kind == "list":
            if state.list_stack:
                state.list_stack.pop()
        elif kind == "list_item":
            state.end_paragraph()
        elif kind == "blockquote":
            state.in_blockquote_depth = max(0, state.in_blockquote_depth - 1)
        elif kind == "table":
            state.end_table()
        elif kind in ("thead", "tbody"):
            state.table_section = None
        elif kind == "tr":
            state.end_table_row()
        elif kind in ("th", "td"):
            state.end_table_cell()

    # ------------------------------------------------------------------
    def _write_attachments(
        self, document: DocumentT, attachments: list[Attachment], stager: AssetStager
    ) -> None:
        head = document.add_paragraph()
        r = head.add_run("Attachments")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for att in attachments:
            staged = stager.lookup_attachment(att)
            p = document.add_paragraph(style="List Bullet")
            if staged is None:
                run = p.add_run(f"{att.filename}  (missing)")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
            else:
                run = p.add_run(att.filename + "  ")
                _add_hyperlink(p, staged.rel_path, staged.rel_path)


# ----------------------------------------------------------------------
# DOCX rendering state machine
# ----------------------------------------------------------------------
class _DocxState:
    def __init__(self, document: DocumentT):
        self.document = document
        self.paragraph = None
        self.list_stack: list[tuple[str, int]] = []  # (kind, depth)
        self.in_blockquote_depth = 0
        self.heading_level: int | None = None
        # Table state
        self.table = None
        self.table_section = None
        self.row_cells: list = []
        self.cell_index = 0
        self.in_header_row = False
        self.in_table_cell = False

    def begin_paragraph(self) -> None:
        if self.table is not None and self.in_table_cell:
            # Inside a table cell we already have a paragraph from add_paragraph
            return
        self.paragraph = self.document.add_paragraph()
        if self.in_blockquote_depth:
            self.paragraph.paragraph_format.left_indent = Inches(
                0.3 * self.in_blockquote_depth
            )
            _add_left_border(self.paragraph, "888888", 12)

    def begin_heading(self, level: int) -> None:
        self.heading_level = level
        self.paragraph = self.document.add_heading(level=min(max(level, 1), 6))
        # Clear any default text
        for r in list(self.paragraph.runs):
            r.text = ""

    def end_paragraph(self) -> None:
        self.paragraph = None
        self.heading_level = None

    def begin_list_item(self) -> None:
        style = "List Number" if (self.list_stack and self.list_stack[-1][0] == "ordered") else "List Bullet"
        depth = len(self.list_stack) - 1
        if depth >= 1 and style == "List Bullet":
            style = "List Bullet 2"
        if depth >= 1 and style == "List Number":
            style = "List Number 2"
        try:
            self.paragraph = self.document.add_paragraph(style=style)
        except KeyError:
            self.paragraph = self.document.add_paragraph()
            self.paragraph.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
            self.paragraph.add_run("• ")
        if self.in_blockquote_depth:
            self.paragraph.paragraph_format.left_indent = Inches(
                0.3 * self.in_blockquote_depth + 0.25 * (depth + 1)
            )

    def add_run(self, run_ev: mdw.Run) -> None:
        if self.paragraph is None:
            self.begin_paragraph()
        text = run_ev.text
        if not text:
            return
        if run_ev.href:
            _add_hyperlink(self.paragraph, text, run_ev.href)
            return
        r = self.paragraph.add_run(text)
        if run_ev.bold:
            r.bold = True
        if run_ev.italic:
            r.italic = True
        if run_ev.code:
            r.font.name = "Consolas"
            _set_complex_font(r, "Consolas")
            r.font.size = Pt(10)
            _shade_run(r, "EFEFEF")

    def hard_break(self) -> None:
        if self.paragraph is None:
            self.begin_paragraph()
        run = self.paragraph.add_run()
        run.add_break(WD_BREAK.LINE)

    # ---------- tables ----------
    def begin_table(self) -> None:
        self.table = self.document.add_table(rows=0, cols=0)
        try:
            self.table.style = "Light Grid Accent 1"
        except KeyError:
            try:
                self.table.style = "Table Grid"
            except KeyError:
                pass

    def begin_table_row(self) -> None:
        if self.table is None:
            return
        self.row_cells = []
        self.cell_index = 0
        self.in_header_row = self.table_section == "thead"

    def begin_table_cell(self, header: bool) -> None:
        if self.table is None:
            return
        # Grow columns as needed
        needed_cols = self.cell_index + 1
        while len(self.table.columns) < needed_cols:
            self.table.add_column(Inches(_PAGE_WIDTH_INCHES / max(needed_cols, 1)))
        if not self.row_cells:
            # Add a new row
            row = self.table.add_row()
            self.row_cells = list(row.cells)
        while len(self.row_cells) <= self.cell_index:
            # Defensive: ensure we have a cell to write into
            self.row_cells.append(self.table.rows[-1].cells[-1])
        cell = self.row_cells[self.cell_index]
        cell.paragraphs[0].clear()
        self.paragraph = cell.paragraphs[0]
        if header:
            for r in self.paragraph.runs:
                r.bold = True
        self.in_table_cell = True

    def end_table_cell(self) -> None:
        if self.in_header_row and self.paragraph is not None:
            for r in self.paragraph.runs:
                r.bold = True
        self.cell_index += 1
        self.in_table_cell = False
        self.paragraph = None

    def end_table_row(self) -> None:
        self.row_cells = []
        self.cell_index = 0
        self.in_header_row = False

    def end_table(self) -> None:
        self.table = None
        self.table_section = None
        self.row_cells = []
        self.cell_index = 0
        self.in_header_row = False

    def finish(self) -> None:
        self.paragraph = None


# ----------------------------------------------------------------------
# DOCX low-level XML helpers
# ----------------------------------------------------------------------
def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _shade_run(run, hex_color: str) -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _add_bottom_border(paragraph, color: str = "CCCCCC", size_eighth_pts: int = 4) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighth_pts))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_left_border(paragraph, color: str = "888888", size_eighth_pts: int = 12) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighth_pts))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_complex_font(run, font_name: str) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a hyperlink run to ``paragraph``. ``url`` may be relative."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _fit_image_width(path: Path, max_inches: float) -> float:
    """Return a width in inches that doesn't exceed ``max_inches`` or
    the image's native size (so small images stay small)."""
    try:
        with Image.open(path) as im:
            w_px, h_px = im.size
            # Assume 96 DPI for screenshots/screen content
            native_w = w_px / 96.0
            return min(max_inches, max(1.0, native_w))
    except Exception:  # noqa: BLE001
        return max_inches
